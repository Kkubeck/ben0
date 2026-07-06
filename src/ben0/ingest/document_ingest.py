"""Ingest text documents into Document and SourceChunk records."""

from __future__ import annotations

import logging
import platform
import re
import subprocess
import tempfile
from pathlib import Path

import docx
import fitz  # pymupdf

from ben0.db.models import Document, SourceChunk
from ben0.db.session import get_session

CHUNK_SIZE = 500


def _classify(filepath: Path) -> str:
    p = str(filepath)
    if "card" in p:
        return "accession_card"
    if "polic" in p:
        return "policy"
    return "other"


def _extract_linked_accession(text: str) -> str | None:
    m = re.search(r"Acc(?:ession)?\.?\s*([\w\-\.]+)", text, re.I)
    return m.group(1) if m else None


def _chunk_text(text: str, size: int = CHUNK_SIZE) -> list[tuple[int, int, str]]:
    """Split text into chunks of ~size chars, splitting on whitespace."""
    chunks = []
    start = 0
    total = len(text)
    while start < total:
        end = min(start + size, total)
        if end < total:
            # Extend to next whitespace boundary
            while end < total and text[end] not in " \n\t":
                end += 1
        chunk = text[start:end].strip()
        if chunk:
            chunks.append((start, end, chunk))
        start = end
    return chunks


def _extract_pdf_text(filepath: Path) -> str | None:
    """Extract text from PDF file page by page."""
    try:
        doc = fitz.open(filepath)
        full_text = ""
        for page_num in range(doc.page_count):
            page = doc.load_page(page_num)
            text = page.get_text()
            full_text += text + "\n"
        doc.close()

        # Return None if no extractable text found
        if not full_text.strip():
            return None
        return full_text.strip()
    except Exception as e:
        logging.warning(f"Failed to extract text from PDF {filepath}: {e}")
        return None


def _extract_docx_text(filepath: Path) -> str | None:
    try:
        doc = docx.Document(filepath)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        if not paragraphs:
            return None
        return "\n".join(paragraphs)
    except Exception as e:
        logging.warning(f"Failed to extract text from DOCX {filepath}: {e}")
        return None


def _extract_doc_text(filepath: Path) -> str | None:
    """Extract text from legacy .doc files using platform tools."""
    try:
        if platform.system() == "Darwin":
            result = subprocess.run(
                ["textutil", "-convert", "txt", "-stdout", str(filepath)],
                capture_output=True, text=True, timeout=30,
            )
            if result.returncode == 0 and result.stdout.strip():
                return result.stdout.strip()
        with tempfile.TemporaryDirectory() as tmpdir:
            result = subprocess.run(
                ["libreoffice", "--headless", "--convert-to", "txt:Text",
                 "--outdir", tmpdir, str(filepath)],
                capture_output=True, text=True, timeout=60,
            )
            if result.returncode == 0:
                txt_path = Path(tmpdir) / (filepath.stem + ".txt")
                if txt_path.exists():
                    text = txt_path.read_text(encoding="utf-8", errors="replace")
                    if text.strip():
                        return text.strip()
    except FileNotFoundError:
        logging.warning(
            f"No .doc converter found for {filepath}. "
            "Install LibreOffice or run on macOS (textutil)."
        )
    except Exception as e:
        logging.warning(f"Failed to extract text from DOC {filepath}: {e}")
    return None


def ingest_documents(doc_dir: Path, db_url: str | None = None) -> dict[str, int]:
    counts = {"documents": 0, "chunks": 0, "skipped": 0}
    session = get_session(db_url)
    try:
        txt_files = sorted(doc_dir.rglob("*.txt"))
        pdf_files = sorted(doc_dir.rglob("*.pdf"))
        docx_files = sorted(doc_dir.rglob("*.docx"))
        doc_files = sorted(doc_dir.rglob("*.doc"))
        doc_files = [f for f in doc_files if f.suffix.lower() == ".doc"]
        all_files = txt_files + pdf_files + docx_files + doc_files

        for filepath in all_files:
            rel = str(filepath.relative_to(doc_dir))
            # Skip if already ingested
            existing = session.query(Document).filter_by(filename=rel).first()
            if existing:
                counts["skipped"] += 1
                continue

            suffix = filepath.suffix.lower()
            if suffix == ".pdf":
                text = _extract_pdf_text(filepath)
                if text is None:
                    logging.warning(f"Skipping PDF with no extractable text: {filepath}")
                    counts["skipped"] += 1
                    continue
            elif suffix == ".docx":
                text = _extract_docx_text(filepath)
                if text is None:
                    logging.warning(f"Skipping DOCX with no extractable text: {filepath}")
                    counts["skipped"] += 1
                    continue
            elif suffix == ".doc":
                text = _extract_doc_text(filepath)
                if text is None:
                    logging.warning(f"Skipping DOC (no converter or no text): {filepath}")
                    counts["skipped"] += 1
                    continue
            else:
                text = filepath.read_text(encoding="utf-8", errors="replace")

            doc_type = _classify(filepath)
            linked_acc = _extract_linked_accession(text) if doc_type == "accession_card" else None
            title = filepath.stem.replace("_", " ").title()

            doc = Document(
                filename=rel,
                document_type=doc_type,
                title=title,
                full_text=text,
                character_count=len(text),
                linked_accession_number=linked_acc,
                sensitivity_level="internal",
                source_file=str(filepath),
            )
            session.add(doc)
            session.flush()

            for idx, (char_start, char_end, chunk_text) in enumerate(
                _chunk_text(text, CHUNK_SIZE)
            ):
                chunk = SourceChunk(
                    document_id=doc.id,
                    chunk_index=idx,
                    chunk_text=chunk_text,
                    char_start=char_start,
                    char_end=char_end,
                )
                session.add(chunk)
                counts["chunks"] += 1

            counts["documents"] += 1

        session.commit()
    except Exception:
        session.rollback()
        raise
    finally:
        session.close()

    return counts
